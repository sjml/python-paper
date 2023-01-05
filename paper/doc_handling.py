import os
import shutil
import subprocess
from datetime import datetime
import zipfile
import tempfile

import typer
from docx import Document
from docx import enum
from docx.shared import Pt
from docx.oxml.ns import qn
from PyPDF2 import PdfFileReader, PdfFileWriter

from . import LIB_NAME, LIB_VERSION_STR
from .shared import PAPER_STATE
from .util import get_date_string


def generate_title_page_string(meta: dict):
    title_string = "::: title-page\n\n"

    if "title" in meta["data"] or "subtitle" in meta["data"]:
        title_string += '::: {custom-style="Title"}\n'
        if "title" in meta["data"]:
            title_string += meta["data"]["title"]
            if "subtitle" in meta["data"]:
                title_string += ":\\\n"
            else:
                title_string += "\n"
        if "subtitle" in meta["data"]:
            title_string += meta["data"]["subtitle"]
            title_string += "\n"
        title_string += ":::\n"

    title_string += '::: {custom-style="Author"}\nby\n:::\n'
    title_string += '::: {custom-style="Author"}\n'
    if "author" in meta["data"]:
        title_string += meta["data"]["author"]
        title_string += "\n"
    title_string += ":::\n"

    info_string = ""
    if "professor" in meta["data"]:
        info_string += f"{meta['data']['professor']}\\\n"
    if "class_mnemonic" in meta["data"]:
        info_string += meta["data"]["class_mnemonic"]
    if "class_mnemonic" in meta["data"] and "class_name" in meta["data"]:
        info_string += " --- "
    if "class_name" in meta["data"]:
        info_string += meta["data"]["class_name"]

    info_string += f"\\\n{get_date_string()}"

    title_string += f'::: {{custom-style="Author"}}\n{info_string}\n:::\n\n'

    title_string += ":::\n"
    return title_string


def package(filename: str, meta: dict):
    if PAPER_STATE["verbose"]:
        typer.echo("Packaging docx...")
    doc = Document(filename)

    # change fonts if we were asked to
    if "base_font_override" in meta and meta["base_font_override"] != None:
        if PAPER_STATE["verbose"]:
            typer.echo(f"Changing base font to {meta['base_font_override']}...")
        doc.styles["Normal"].font.name = meta["base_font_override"]
    if "mono_font_override" in meta and meta["mono_font_override"] != None:
        if PAPER_STATE["verbose"]:
            typer.echo(f"Changing mono font to {meta['mono_font_override']}...")
        doc.styles["Verbatim Char"].font.name = meta["mono_font_override"]

    # set metadata
    if PAPER_STATE["verbose"]:
        typer.echo("Fixing docx metadata...")
    doc.core_properties.title = meta["data"]["title"]
    doc.core_properties.author = meta["data"]["author"]
    doc.core_properties.last_modified_by = meta["data"]["author"]
    if PAPER_STATE["docx"]["revision"] <= 0:
        doc.core_properties.revision = max(1, int(subprocess.check_output(["git", "rev-list", "--all", "--count"])) - 1)
    else:
        doc.core_properties.revision = int(PAPER_STATE["docx"]["revision"])

    # check "Total Row" box on tables
    if len(doc.tables) > 0:
        if PAPER_STATE["verbose"]:
            typer.echo("Fixing table formatting...")
    for t in doc.tables:
        tblLook = t._tblPr.first_child_found_in("w:tblLook")
        tblLook.set(qn("w:lastRow"), "1")

    if PAPER_STATE["verbose"]:
        typer.echo(f"Writing final docx to {filename}...")
    doc.save(filename)

    if "SOURCE_DATE_EPOCH" in os.environ:
        if PAPER_STATE["verbose"]:
            typer.echo(f"Correcting interior timestamps on {filename}...")

        epoch = float(os.environ["SOURCE_DATE_EPOCH"])
        with tempfile.TemporaryDirectory() as tmpdir:
            zf = zipfile.ZipFile(filename)
            zf.extractall(tmpdir)
            zf.close()
            os.unlink(filename)
            nzf = zipfile.ZipFile(filename, "w", compression=zipfile.ZIP_DEFLATED)
            for dirpath, _, files in os.walk(tmpdir):
                for f in files:
                    fpath = os.path.join(dirpath, f)
                    os.utime(fpath, (epoch, epoch))
                    nzf.write(fpath, fpath[len(tmpdir) :])
            nzf.close()


def make_pdf(filename: str, meta: dict):
    if PAPER_STATE["verbose"]:
        typer.echo("Generating PDF...")
    filepath = os.path.abspath(filename)
    base = os.path.splitext(filepath)[0]
    pdf_filepath = base + ".pdf"
    if os.path.exists(pdf_filepath):
        os.unlink(pdf_filepath)

    outpath = os.path.expanduser(
        f"~/Library/Containers/com.microsoft.Word/Data/Documents/{os.path.basename(pdf_filepath)}"
    )

    applescript = f"""
    tell application "System Events" to set wordIsRunning to exists (processes where name is "Microsoft Word")

    set outpath to "{outpath[1:].replace('/', ':')}"
    tell application id "com.microsoft.Word"
        activate
        open file "{filepath[1:].replace('/', ':')}"
        repeat while not (active document is not missing value)
            delay 0.5
        end repeat
        set activeDoc to active document
        save as activeDoc file name outpath file format format PDF
        close activeDoc saving no
        if not wordIsRunning then
            quit
        end if
    end tell
    """
    subprocess.run(["osascript", "-"], input=applescript.encode("utf-8"))

    shutil.move(outpath, pdf_filepath)

    if PAPER_STATE["verbose"]:
        typer.echo("Fixing PDF metadata...")
    reader = PdfFileReader(pdf_filepath)
    writer = PdfFileWriter()
    writer.appendPagesFromReader(reader)
    pdf_date = datetime.utcnow().strftime("%Y%m%d%H%MZ00'00'")
    writer.addMetadata(
        {
            "/Author": meta["data"]["author"],
            "/Title": meta["data"]["title"],
            "/Producer": f"{LIB_NAME} {LIB_VERSION_STR}",
            "/CreationDate": f"D:{pdf_date}",
        }
    )

    if PAPER_STATE["verbose"]:
        typer.echo(f"Writing final PDF to {pdf_filepath}...")
    with open(pdf_filepath, "wb") as pdfout:
        writer.write(pdfout)
